"""
Resume Formatter — generates Word doc matching the NexTurn template exactly.
Layout: two-column table (main left, grey sidebar right).
Page 1: Name, Role, Objective, Work Experience (left) | Edu, Skills (right sidebar)
Page 2: Technologies Used + SKILLS breakdown

Uses python-docx only (no external template engine needed).
"""

import io
import json
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from utils.groq_client import create_groq_completion


# ── Colours (from ELIZABETH_CARTER reference) ─────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)
MID_GREY   = RGBColor(0x44, 0x44, 0x44)
BLUE_DATE  = RGBColor(0x2E, 0x74, 0xB5)
LIGHT_GREY = RGBColor(0x55, 0x55, 0x55)
SIDEBAR_BG = "B8C4CA"   # hex for shading XML
RED        = RGBColor(0xFF, 0x00, 0x00)


# ── LLM extraction ────────────────────────────────────────────────────────────

def extract_detailed_resume_data(client, resume_text: str, candidate_meta: dict) -> dict:
    fallback_client = st.session_state.get('fallback_client')

    prompt = f"""You are an expert resume parser. Extract ALL details from the resume below into structured JSON.

RULES:
- Extract exact text from the resume. Do not fabricate.
- For work_experience bullets, extract actual bullet text verbatim.
- Return ONLY this JSON (no markdown):
{{
  "NAME": "Full Name",
  "ROLE": "Job Title",
  "PROFESSIONAL_SUMMARY": "2-3 sentence summary",
  "experience_years": "X+",
  "COMPANY_NAME": "Most recent company",
  "LOCATION": "City, Country",
  "START_DATE": "Mon YYYY",
  "END_DATE": "Mon YYYY or Present",
  "PROJECT1_NAME": "Primary project title",
  "ABOUT_PROJECT_BULLET_1": "bullet 1",
  "ABOUT_PROJECT_BULLET_2": "bullet 2",
  "ABOUT_PROJECT_BULLET_3": "bullet 3",
  "ABOUT_PROJECT_BULLET_4": "bullet 4",
  "ABOUT_PROJECT_BULLET_5": "bullet 5",
  "ABOUT_PROJECT_BULLET_6": "bullet 6",
  "PROJECT2_NAME": "Second project title or empty",
  "PROJECT2_BULLET_1": "bullet 1",
  "PROJECT2_BULLET_2": "bullet 2",
  "PROJECT2_BULLET_3": "bullet 3",
  "PROJECT2_BULLET_4": "bullet 4",
  "PROJECT2_BULLET_5": "bullet 5",
  "PROJECT2_BULLET_6": "bullet 6",
  "TECHNOLOGIES_USED": "Comma separated tech list",
  "HIGHEST_EDUCATION": "Degree name",
  "COLLEGE_NAME": "University name",
  "EDUCATION_DATES": "Year range",
  "tech_stack": ["Skill1", "Skill2", "Skill3", "Skill4", "Skill5", "Skill6", "Skill7", "Skill8", "Skill9", "Skill10", "Skill11", "Skill12"],
  "BACKEND_LANGUAGES": "e.g. Python, Java",
  "CONTAINERS_AND_ORCHESTRATION": "e.g. Docker, Kubernetes",
  "DATABASES": "e.g. PostgreSQL, MongoDB",
  "OPERATING_SYSTEMS": "e.g. Linux, Windows",
  "VERSION_CONTROL_TOOLS": "e.g. Git, GitHub",
  "TESTING_TOOLS": "e.g. PyTest, Selenium"
}}

Resume:
{resume_text[:7000]}"""

    try:
        response = create_groq_completion(
            client, fallback_client,
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a precise resume parser. Return only valid JSON with no markdown."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=3000
        )
        content = response.choices[0].message.content.strip()
        j_start = content.find('{')
        j_end   = content.rfind('}') + 1
        if j_start != -1 and j_end > j_start:
            data = json.loads(content[j_start:j_end])
            # Fallback from meta for critical fields
            if not data.get('NAME'):        data['NAME']  = candidate_meta.get('name', '')
            if not data.get('ROLE'):        data['ROLE']  = candidate_meta.get('current_role', '')
            if not data.get('experience_years'):
                data['experience_years'] = str(candidate_meta.get('experience_years', ''))
            return data
    except Exception as e:
        st.warning(f"Could not fully parse resume details: {e}")

    # Fallback
    return {
        "NAME": candidate_meta.get('name', ''),
        "ROLE": candidate_meta.get('current_role', ''),
        "PROFESSIONAL_SUMMARY": candidate_meta.get('objective', ''),
        "experience_years": str(candidate_meta.get('experience_years', '')),
        "COMPANY_NAME": "", "LOCATION": "",
        "START_DATE": "", "END_DATE": "",
        "PROJECT1_NAME": "",
        **{f"ABOUT_PROJECT_BULLET_{i}": "" for i in range(1, 7)},
        "PROJECT2_NAME": "",
        **{f"PROJECT2_BULLET_{i}": "" for i in range(1, 7)},
        "TECHNOLOGIES_USED": "",
        "HIGHEST_EDUCATION": candidate_meta.get('education', ''),
        "COLLEGE_NAME": "", "EDUCATION_DATES": "",
        "tech_stack": candidate_meta.get('tech_stack', []),
        "BACKEND_LANGUAGES": "", "CONTAINERS_AND_ORCHESTRATION": "",
        "DATABASES": "", "OPERATING_SYSTEMS": "",
        "VERSION_CONTROL_TOOLS": "", "TESTING_TOOLS": "",
    }


def check_template_completeness(data: dict) -> dict:
    warnings = []
    if not data.get('NAME'):              warnings.append("Candidate name not found")
    if not data.get('ROLE'):              warnings.append("Job title not found")
    if not data.get('experience_years'):  warnings.append("Years of experience not specified")
    if not data.get('COMPANY_NAME'):      warnings.append("Company name missing")
    if not data.get('PROJECT1_NAME'):     warnings.append("Primary project not found")
    if not data.get('PROFESSIONAL_SUMMARY'): warnings.append("Professional summary missing")
    if not data.get('HIGHEST_EDUCATION'): warnings.append("Education details not found")
    critical = not data.get('NAME') or not data.get('COMPANY_NAME')
    return {'complete': len(warnings) == 0, 'warnings': warnings, 'has_critical_gaps': critical}


# ── python-docx helpers ───────────────────────────────────────────────────────

def _sf(run, size=11, bold=False, italic=False, color=None, name='Arial'):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _para(container, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=2):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    return p


def _add_bottom_border(para, color_hex='2E74B5', sz='6'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _section_heading(container, text, color=DARK_BLUE):
    p = _para(container, sb=8, sa=2)
    r = p.add_run(text.upper())
    _sf(r, size=11, bold=True, color=color)
    _add_bottom_border(p, color_hex='1F3864')
    return p


def _shade_cell(cell, fill_hex=SIDEBAR_BG):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def _remove_cell_border(table):
    """Remove all borders from a table."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _val(d, key, fallback=''):
    v = d.get(key, fallback)
    return str(v).strip() if v else fallback


def _bullet_para(container, text, size=11):
    p = container.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    _sf(r, size=size)
    return p


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_resume_docx(data: dict) -> bytes | None:
    try:
        doc = Document()

        # ── Page layout: Letter size, standard margins ──────────────────────
        for sec in doc.sections:
            sec.page_width    = Inches(8.5)
            sec.page_height   = Inches(11.0)
            sec.top_margin    = Inches(0.75)
            sec.bottom_margin = Inches(0.75)
            sec.left_margin   = Inches(0.75)
            sec.right_margin  = Inches(0.75)

        # ── Default style ─────────────────────────────────────────────────────
        ns = doc.styles['Normal']
        ns.font.name = 'Arial'
        ns.font.size = Pt(11)

        name    = _val(data, 'NAME',  'Candidate')
        role    = _val(data, 'ROLE',  'Professional')
        exp_yrs = _val(data, 'experience_years', '')
        summary = _val(data, 'PROFESSIONAL_SUMMARY')

        # ─────────────────────────────────────────────────────────────────────
        # HEADER: Name (centred, dark blue large) + Role (centred, grey)
        # ─────────────────────────────────────────────────────────────────────
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(0)
        p_name.paragraph_format.space_after  = Pt(2)
        r = p_name.add_run(f'\t{name.upper()}\t ')
        _sf(r, size=20, bold=True, color=DARK_BLUE)

        p_role = doc.add_paragraph()
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_role.paragraph_format.space_before = Pt(0)
        p_role.paragraph_format.space_after  = Pt(0)
        r = p_role.add_run(f'\t{role.upper()}\t ')
        _sf(r, size=11, color=MID_GREY)

        # Divider line under header
        p_div = _para(doc, sb=4, sa=6)
        _add_bottom_border(p_div, color_hex='2E74B5', sz='12')

        # ─────────────────────────────────────────────────────────────────────
        # TWO-COLUMN TABLE: left=main content, right=grey sidebar
        # ─────────────────────────────────────────────────────────────────────
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _remove_cell_border(table)

        # Column widths: 65/35 split within 7 inch usable width
        col_widths = [Inches(4.55), Inches(2.45)]
        for i, col in enumerate(table.columns):
            for cell in col.cells:
                cell.width = col_widths[i]

        left_cell  = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        left_cell.vertical_alignment  = WD_ALIGN_VERTICAL.TOP
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade_cell(right_cell)

        lc = left_cell    # shorthand
        rc = right_cell

        # ── LEFT COLUMN ───────────────────────────────────────────────────────

        # OBJECTIVE / Summary
        _section_heading(lc, 'OBJECTIVE')
        if summary:
            p = _para(lc, sb=0, sa=4)
            r = p.add_run(summary)
            _sf(r, size=11)

        # WORK EXPERIENCE heading
        exp_label = f'WORK EXPERIENCE ({exp_yrs}YRS)' if exp_yrs else 'WORK EXPERIENCE'
        _section_heading(lc, exp_label)

        # Company + location
        company  = _val(data, 'COMPANY_NAME')
        location = _val(data, 'LOCATION')
        if company:
            p = _para(lc, sb=6, sa=1)
            r1 = p.add_run(company)
            _sf(r1, bold=True)
            if location:
                r2 = p.add_run(f'   {location}')
                _sf(r2, color=LIGHT_GREY)

        # Dates
        start = _val(data, 'START_DATE')
        end   = _val(data, 'END_DATE')
        if start or end:
            p = _para(lc, sb=0, sa=1)
            r = p.add_run(f'{start}-{end}' if start and end else start or end)
            _sf(r, bold=True, color=BLUE_DATE)

        # Project 1
        p1_name = _val(data, 'PROJECT1_NAME')
        if p1_name:
            p = _para(lc, sb=2, sa=1)
            r = p.add_run(f'Project: {p1_name}')
            _sf(r, bold=True)

        for i in range(1, 7):
            bullet = _val(data, f'ABOUT_PROJECT_BULLET_{i}')
            if bullet:
                _bullet_para(lc, bullet)

        # Project 2
        p2_name = _val(data, 'PROJECT2_NAME')
        if p2_name:
            p = _para(lc, sb=8, sa=1)
            r = p.add_run(f'Projects: {p2_name}')
            _sf(r, bold=True)

            for i in range(1, 7):
                bullet = _val(data, f'PROJECT2_BULLET_{i}')
                if bullet:
                    _bullet_para(lc, bullet)

        # Technologies
        tech_used = _val(data, 'TECHNOLOGIES_USED')
        if tech_used:
            p = _para(lc, sb=6, sa=4)
            r1 = p.add_run('     Technologies: ')
            _sf(r1, bold=True)
            r2 = p.add_run(tech_used)
            _sf(r2)

        # ── RIGHT SIDEBAR ─────────────────────────────────────────────────────

        # Remove the default empty paragraph that python-docx adds to cells
        for para in rc.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # ACADEMIC QUALIFICATIONS
        _section_heading(rc, 'ACADEMIC QUALIFICATIONS', color=DARK_BLUE)

        highest_edu = _val(data, 'HIGHEST_EDUCATION')
        college     = _val(data, 'COLLEGE_NAME')
        edu_dates   = _val(data, 'EDUCATION_DATES')

        for val in [highest_edu, college, edu_dates]:
            if val:
                p = rc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                r = p.add_run(val)
                _sf(r, size=11)

        # KEY SKILLS
        p_sk = rc.add_paragraph()
        p_sk.paragraph_format.space_before = Pt(12)
        p_sk.paragraph_format.space_after  = Pt(2)
        r = p_sk.add_run('KEY SKILLS')
        _sf(r, size=11, bold=True, color=DARK_BLUE)

        tech_stack = data.get('tech_stack') or []
        if isinstance(tech_stack, str):
            tech_stack = [s.strip() for s in tech_stack.split(',') if s.strip()]

        for skill in tech_stack[:12]:
            if skill:
                _bullet_para(rc, str(skill), size=10)

        # ─────────────────────────────────────────────────────────────────────
        # PAGE 2: SKILLS breakdown
        # ─────────────────────────────────────────────────────────────────────
        doc.add_page_break()

        skills_sections = [
            ('Backend Languages',          'BACKEND_LANGUAGES'),
            ('Containers & Orchestration', 'CONTAINERS_AND_ORCHESTRATION'),
            ('Databases',                  'DATABASES'),
            ('Operating Systems',          'OPERATING_SYSTEMS'),
            ('Version Control & Tools',    'VERSION_CONTROL_TOOLS'),
            ('Testing',                    'TESTING_TOOLS'),
        ]

        has_skills_page = any(_val(data, k) for _, k in skills_sections)

        if has_skills_page:
            _section_heading(doc, 'SKILLS')

            for label, key in skills_sections:
                val = _val(data, key)
                if val:
                    p = _para(doc, sb=4, sa=0)
                    r = p.add_run(f'{label}:')
                    _sf(r, bold=True)
                    p2 = _para(doc, sb=0, sa=4)
                    r2 = p2.add_run(val)
                    _sf(r2)

        # ── Save ──────────────────────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    except Exception as e:
        st.error(f"Could not create document: {e}")
        return None